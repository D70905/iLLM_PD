"""Add Table 2 (component ablation) to the manuscript after the ablation paragraph."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement as OE

doc = Document(r'C:/Users/Ivy/Desktop/NC一审意见回复/文本/manuscript/iLLM-PD_重写_0608.docx')

# Find the ablation paragraph
ablation_idx = None
for i, p in enumerate(doc.paragraphs):
    if 'We removed the Generator' in p.text and 'four variants' in p.text:
        ablation_idx = i
        break

if ablation_idx is None:
    raise SystemExit("ERROR: ablation paragraph not found")

print(f'Ablation paragraph at index {ablation_idx}')
body = doc.element.body

# Find the paragraph element in body
body_idx = None
for i, child in enumerate(body):
    if child.tag == qn('w:p'):
        texts = child.findall('.//' + qn('w:t'))
        combined = ''.join(t.text or '' for t in texts)
        if 'We removed the Generator' in combined:
            body_idx = i
            break

if body_idx is None:
    raise SystemExit("ERROR: paragraph not found in body")

def add_paragraph_to_body(body, insert_idx, text, font_size, bold=False, italic=False, align='left'):
    """Add a simple paragraph at insert_idx in the body."""
    p_el = OE('w:p')
    pPr = OE('w:pPr')
    jc = OE('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    p_el.append(pPr)

    r_el = OE('w:r')
    rPr = OE('w:rPr')
    sz = OE('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))  # half-points
    rPr.append(sz)
    if bold:
        b = OE('w:b')
        rPr.append(b)
    if italic:
        i_el = OE('w:i')
        rPr.append(i_el)
    r_el.append(rPr)

    t_el = OE('w:t')
    t_el.set(qn('xml:space'), 'preserve')
    t_el.text = text
    r_el.append(t_el)
    p_el.append(r_el)

    body.insert(insert_idx, p_el)
    return insert_idx + 1

# Insert table caption and sub-caption
offset = 1
body_idx = add_paragraph_to_body(body, body_idx + offset,
    'Table 2 | Component ablation (4 variants × 2 pavement types × 3 seeds; 1,000 timesteps each; deterministic inference on LTPP sections).',
    9, bold=True)
offset += 1

body_idx = add_paragraph_to_body(body, body_idx + offset,
    'DSR = design safety rate (weakest-link); SCR = episode specification-compliance rate; Cost = delivered construction cost (USD m⁻²). Mean ± s.d. over 6 sections × 3 seeds. Δ cost vs Full.',
    8)
offset += 1

# Build the table
table_data = [
    ['Variant', 'Type', 'DSR', 'SCR', 'Cost (USD/m²)', 'Δ cost'],
    ['Full', 'flexible', '1.000 ± 0.000', '0.997 ± 0.011', '37.6 ± 0.4', '—'],
    ['No Generator', 'flexible', '1.000 ± 0.000', '0.847 ± 0.202', '38.2 ± 0.6', '+1.6%'],
    ['No RAG', 'flexible', '1.000 ± 0.000', '0.937 ± 0.106', '40.8 ± 1.8', '+8.5%'],
    ['No Guard', 'flexible', '1.000 ± 0.000', '0.659 ± 0.321', '41.6 ± 0.5', '+10.6%'],
    ['Full', 'semi-rigid', '1.000 ± 0.000', '1.000 ± 0.000', '54.7 ± 1.5', '—'],
    ['No Generator', 'semi-rigid', '1.000 ± 0.000', '1.000 ± 0.000', '55.6 ± 1.2', '+1.6%'],
    ['No RAG', 'semi-rigid', '1.000 ± 0.000', '1.000 ± 0.000', '54.8 ± 1.5', '+0.2%'],
    ['No Guard', 'semi-rigid', '1.000 ± 0.000', '1.000 ± 0.000', '52.8 ± 2.5', '−3.5%'],
]

tbl = OE('w:tbl')

# Table properties
tblPr = OE('w:tblPr')
tblW = OE('w:tblW'); tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa')
tblPr.append(tblW)
tblBorders = OE('w:tblBorders')
for bn in ['top','left','bottom','right','insideH','insideV']:
    b = OE('w:'+bn); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),'555555')
    tblBorders.append(b)
tblPr.append(tblBorders)
tbl.append(tblPr)

tblGrid = OE('w:tblGrid')
col_w = [1800, 1400, 1700, 1700, 1700, 1000]
for w in col_w:
    gc = OE('w:gridCol'); gc.set(qn('w:w'), str(w)); tblGrid.append(gc)
tbl.append(tblGrid)

for ri, rd in enumerate(table_data):
    tr = OE('w:tr')
    for ci, ct in enumerate(rd):
        tc = OE('w:tc')
        tcPr = OE('w:tcPr')
        tcW = OE('w:tcW'); tcW.set(qn('w:w'), str(col_w[ci])); tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        if ri == 0:
            shd = OE('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'E0E0E0')
            tcPr.append(shd)
        tc.append(tcPr)

        p_el = OE('w:p')
        pPr = OE('w:pPr')
        jc = OE('w:jc'); jc.set(qn('w:val'), 'left' if ci < 2 else 'center')
        pPr.append(jc)
        p_el.append(pPr)

        r_el = OE('w:r')
        rPr = OE('w:rPr')
        rFonts = OE('w:rFonts'); rFonts.set(qn('w:ascii'),'Arial'); rFonts.set(qn('w:hAnsi'),'Arial')
        rPr.append(rFonts)
        sz = OE('w:sz'); sz.set(qn('w:val'), '16' if ri > 0 else '18')
        rPr.append(sz)
        if ri == 0:
            b_el = OE('w:b'); rPr.append(b_el)
        r_el.append(rPr)
        t_el = OE('w:t'); t_el.set(qn('xml:space'),'preserve'); t_el.text = str(ct)
        r_el.append(t_el)
        p_el.append(r_el)
        tc.append(p_el)
        tr.append(tc)
    tbl.append(tr)

body.insert(body_idx + offset, tbl)
offset += 1

# Footnote
body_idx = add_paragraph_to_body(body, body_idx + offset,
    'The semi-rigid cost differences are within seed noise (± 2.5) and not statistically distinguishable from Full. The ablation used a reduced 1,000-timestep budget (vs. 2,048 for the main flexible policy); it is a relative comparison across variants, not a re-run of the headline result.',
    7, italic=True)
offset += 1

doc.save(r'C:/Users/Ivy/Desktop/NC一审意见回复/文本/manuscript/iLLM-PD_重写_0608.docx')
print('Saved with Table 2 inserted.')